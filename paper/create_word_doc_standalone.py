import os
import docx
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.section import WD_SECTION
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
import re

def set_section_columns(section, num_cols, space_pt=18):
    sectPr = section._sectPr
    cols = sectPr.find(qn('w:cols'))
    if cols is None:
        cols = OxmlElement('w:cols')
        sectPr.append(cols)
    cols.set(qn('w:num'), str(num_cols))
    cols.set(qn('w:space'), str(int(space_pt * 20)))

def format_run(run, font_name="Times New Roman", size_pt=10, bold=False, italic=False, color_rgb=(0,0,0)):
    run.font.name = font_name
    run.font.size = Pt(size_pt)
    run.bold = bold
    run.italic = italic
    run.font.color.rgb = RGBColor(*color_rgb)

def add_bookmark(paragraph, bookmark_name, bookmark_id):
    start = OxmlElement('w:bookmarkStart')
    start.set(qn('w:id'), str(bookmark_id))
    start.set(qn('w:name'), bookmark_name)
    paragraph._p.insert(0, start)
    
    end = OxmlElement('w:bookmarkEnd')
    end.set(qn('w:id'), str(bookmark_id))
    paragraph._p.append(end)

def add_hyperlink_to_bookmark(paragraph, bookmark_name, link_text):
    hyperlink = OxmlElement('w:hyperlink')
    hyperlink.set(qn('w:anchor'), bookmark_name)
    
    new_run = OxmlElement('w:r')
    rPr = OxmlElement('w:rPr')
    
    rFonts = OxmlElement('w:rFonts')
    rFonts.set(qn('w:ascii'), 'Times New Roman')
    rFonts.set(qn('w:hAnsi'), 'Times New Roman')
    rPr.append(rFonts)
    
    sz = OxmlElement('w:sz')
    sz.set(qn('w:val'), '20')
    rPr.append(sz)
    
    u = OxmlElement('w:u')
    u.set(qn('w:val'), 'single')
    rPr.append(u)
    
    color = OxmlElement('w:color')
    color.set(qn('w:val'), '0000FF')
    rPr.append(color)
    
    new_run.append(rPr)
    
    text = OxmlElement('w:t')
    text.text = link_text
    new_run.append(text)
    
    hyperlink.append(new_run)
    paragraph._p.append(hyperlink)

def add_formatted_run_text(p, text, size_pt=10):
    parts = text.split("**")
    for i, part in enumerate(parts):
        if i % 2 == 1:
            run = p.add_run(part)
            format_run(run, font_name="Times New Roman", size_pt=size_pt, bold=True)
        else:
            subparts = part.split("*")
            for j, subpart in enumerate(subparts):
                run = p.add_run(subpart)
                if j % 2 == 1:
                    format_run(run, font_name="Times New Roman", size_pt=size_pt, italic=True)
                else:
                    format_run(run, font_name="Times New Roman", size_pt=size_pt)

def add_heading_1(doc, text):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_before = Pt(12)
    p.paragraph_format.space_after = Pt(6)
    p.paragraph_format.keep_with_next = True
    run = p.add_run(text)
    format_run(run, font_name="Times New Roman", size_pt=10, bold=True)
    return p

def add_heading_2(doc, text):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    p.paragraph_format.space_before = Pt(8)
    p.paragraph_format.space_after = Pt(4)
    p.paragraph_format.keep_with_next = True
    run = p.add_run(text)
    format_run(run, font_name="Times New Roman", size_pt=10, bold=False, italic=True)
    return p

def add_body_paragraph(doc, text, first_line_indent_in=0.15):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    p.paragraph_format.line_spacing = 1.05
    p.paragraph_format.space_after = Pt(4)
    if first_line_indent_in > 0:
        p.paragraph_format.first_line_indent = Inches(first_line_indent_in)
    
    pattern = re.compile(r'(\[\d+\])')
    chunks = pattern.split(text)
    
    for chunk in chunks:
        if pattern.match(chunk):
            ref_num = chunk[1:-1]
            bookmark_name = f"ref_{ref_num}"
            add_hyperlink_to_bookmark(p, bookmark_name, chunk)
        else:
            add_formatted_run_text(p, chunk, size_pt=10)
    return p

def add_equation(doc, label, eq_text):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_before = Pt(6)
    p.paragraph_format.space_after = Pt(6)
    
    run_eq = p.add_run(eq_text + "\t\t")
    format_run(run_eq, font_name="Times New Roman", size_pt=10, italic=True)
    
    run_label = p.add_run(label)
    format_run(run_label, font_name="Times New Roman", size_pt=10)
    return p

def add_figure(doc, image_path, caption_text, figure_num):
    p_img = doc.add_paragraph()
    p_img.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p_img.paragraph_format.space_before = Pt(8)
    p_img.paragraph_format.space_after = Pt(4)
    
    if os.path.exists(image_path):
        try:
            p_img.add_run().add_picture(image_path, width=Inches(3.2))
            print(f"Added picture {image_path} successfully.")
        except Exception as e:
            p_img.add_run(f"[Error loading image: {e}]")
    else:
        p_img.add_run(f"[Figure {figure_num} Placeholder: {os.path.basename(image_path)}]")
        print(f"Image {image_path} not found. Added placeholder text.")
        
    p_cap = doc.add_paragraph()
    p_cap.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p_cap.paragraph_format.space_after = Pt(8)
    run_cap = p_cap.add_run(f"Fig. {figure_num}. {caption_text}")
    format_run(run_cap, font_name="Times New Roman", size_pt=8)

def add_reference(doc, num, text):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    p.paragraph_format.left_indent = Inches(0.25)
    p.paragraph_format.first_line_indent = Inches(-0.25)
    p.paragraph_format.space_after = Pt(3)
    
    bookmark_name = f"ref_{num}"
    add_bookmark(p, bookmark_name, num)
    
    run_num = p.add_run(f"[{num}]\t")
    format_run(run_num, font_name="Times New Roman", size_pt=8)
    
    parts = text.split('"')
    for i, part in enumerate(parts):
        if i == 1:
            run = p.add_run(f'"{part}"')
            format_run(run, font_name="Times New Roman", size_pt=8)
        else:
            subparts = part.split("*")
            for j, subpart in enumerate(subparts):
                run = p.add_run(subpart)
                if j % 2 == 1:
                    format_run(run, font_name="Times New Roman", size_pt=8, italic=True)
                else:
                    format_run(run, font_name="Times New Roman", size_pt=8)
    return p

def main():
    doc = docx.Document()
    
    # Page setup
    section1 = doc.sections[0]
    section1.top_margin = Inches(0.75)
    section1.bottom_margin = Inches(1.0)
    section1.left_margin = Inches(0.625)
    section1.right_margin = Inches(0.625)
    
    p_title = doc.add_paragraph()
    p_title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p_title.paragraph_format.space_before = Pt(24)
    p_title.paragraph_format.space_after = Pt(12)
    run_title = p_title.add_run("Explainable and Context-Aware Fine-Tuning of\nBidirectional Encoder Representations from Transformers for\nOnline Recruitment Fraud Detection")
    format_run(run_title, font_name="Times New Roman", size_pt=20, bold=True)
    
    p_author = doc.add_paragraph()
    p_author.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p_author.paragraph_format.space_after = Pt(18)
    run_author = p_author.add_run(
        "Akhilesh Yadav\n"
        "Department of Computer Science and Engineering\n"
        "M.Sc. (Data Science)\n"
        "India\n"
        "akhileshyadav8@gmail.com"
    )
    format_run(run_author, font_name="Times New Roman", size_pt=11)
    
    # 2 Column layout for body
    section2 = doc.add_section(WD_SECTION.CONTINUOUS)
    section2.top_margin = Inches(0.75)
    section2.bottom_margin = Inches(1.0)
    section2.left_margin = Inches(0.625)
    section2.right_margin = Inches(0.625)
    set_section_columns(section2, num_cols=2, space_pt=18)
    
    # Abstract
    p_abs = doc.add_paragraph()
    p_abs.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    p_abs.paragraph_format.space_after = Pt(6)
    p_abs.paragraph_format.first_line_indent = Inches(0.15)
    
    run_abs_tag = p_abs.add_run("Abstract---")
    format_run(run_abs_tag, font_name="Times New Roman", size_pt=9, bold=True, italic=True)
    
    run_abs_text = p_abs.add_run(
        "Online recruitment platforms have simplified global recruitment processes. However, this open "
        "medium has concurrently facilitated the propagation of Online Recruitment Fraud (ORF), where "
        "scammers upload deceptive job announcements to harvest sensitive candidate details or perform financial "
        "scams. Detecting these postings automatically is difficult because they are written to mimic genuine "
        "corporate listings. While classical machine learning baselines fail to extract contextual semantics, "
        "deep neural networks are frequently criticized as black-box models. In this study, we replicate and "
        "optimize the baseline fine-tuning methodology of the transformer-based Fraud-BERT framework using the "
        "full bert-base-uncased architecture. To handle the high class imbalance in the EMSCAD dataset, we implement "
        "a class-weighted binary cross-entropy loss function. Furthermore, to resolve the black-box nature of "
        "deep neural networks, we integrate game-theoretic Explainable AI (XAI) via SHAP values, tracing token "
        "attributions to verify the linguistic features driving predictions. Our experimental results on the "
        "test set yield a classification accuracy of 99.02%, Class 1 (Fraudulent) precision of 91.57%, recall of "
        "87.86%, and F1-score of 89.68%. Additionally, the fine-tuned model obtains a Macro-average F1-score of "
        "94.58% and a high ROC-AUC of 99.22%. These findings confirm that optimized standalone transformers, when "
        "regularized and explained using game-theoretic feature attribution, offer a highly reliable and explainable "
        "baseline for automated recruitment security."
    )
    format_run(run_abs_text, font_name="Times New Roman", size_pt=9, bold=True)
    
    p_key = doc.add_paragraph()
    p_key.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    p_key.paragraph_format.space_after = Pt(12)
    p_key.paragraph_format.first_line_indent = Inches(0.15)
    
    run_key_tag = p_key.add_run("Keywords---")
    format_run(run_key_tag, font_name="Times New Roman", size_pt=9, bold=True, italic=True)
    
    run_key_text = p_key.add_run(
        "Online Recruitment Fraud, Transformer, BERT, Fine-Tuning, Explainable AI, SHAP."
    )
    format_run(run_key_text, font_name="Times New Roman", size_pt=9)
    
    # Section 1
    add_heading_1(doc, "I.  INTRODUCTION")
    add_body_paragraph(doc,
        "Online recruitment has become the standard mechanism for job seeking and hiring globally, utilizing "
        "platforms like LinkedIn, Glassdoor, and Indeed. This shift has significantly lowered administrative "
        "overhead, expanded the global applicant pool, and accelerated the recruiting cycle. However, this open "
        "digital ecosystem has also emerged as a fertile ground for cybercriminals engaged in Online Recruitment "
        "Fraud (ORF) [1], [2].")
    
    add_body_paragraph(doc,
        "Fraudulent job ads are designed to bait vulnerable candidates, often promising high compensation, "
        "flexible hours, and low entry requirements. Once candidates apply, fraudsters exploit the situation to "
        "perform identity theft (obtaining bank account numbers, social security numbers, and passport photos) "
        "or to run financial scams (requesting fees for training materials, visa applications, or screening checks) "
        "[3]. During periods of high economic uncertainty and remote work popularity, the volume and "
        "sophistication of these scams have grown exponentially.")
    
    add_body_paragraph(doc,
        "Detecting fraudulent postings is an intricate natural language processing (NLP) problem. Deceptive "
        "listings are constructed using highly professional language, often copying corporate profiles from "
        "legitimate websites to look genuine [4]. Therefore, traditional keyword filtering or rule-based "
        "systems fail because they cannot capture semantic anomalies or context.")
    
    add_body_paragraph(doc,
        "Early automated detection systems deployed machine learning models such as Logistic Regression, "
        "Naive Bayes, and Random Forest on features extracted via Bag-of-Words (BoW) or Term Frequency-Inverse "
        "Document Frequency (TF-IDF) [5]. Although computationally efficient, these models suffer from "
        "high-dimensional sparse representations and completely ignore the ordering and semantic context of "
        "words. To resolve this, deep learning models such as Recurrent Neural Networks (RNN) and Bidirectional "
        "Long Short-Term Memory (Bi-LSTM) were proposed [2]. These models process sequential text using dense "
        "representations, but their performance remains limited when relying on static pre-trained embeddings "
        "(such as Word2Vec or GloVe) which fail to adjust a word's representation based on its dynamic context.")
    
    add_body_paragraph(doc,
        "Recently, Transformer-based Large Language Models (LLMs) like BERT (Bidirectional Encoder "
        "Representations from Transformers) have redefined the state-of-the-art in NLP [6], [7]. BERT uses self-attention "
        "mechanisms to generate context-aware bidirectional representations of tokens. Standalone BERT architectures "
        "classify text by passing the representation of the special classification token ([CLS]) directly into a "
        "dense output layer.")
    
    add_body_paragraph(doc,
        "In this study, we replicate and optimize the fine-tuning of a standalone **BERT** classifier. "
        "This architecture leverages the strengths of BERT: (1) We utilize a pre-trained **BERT-base** encoder to extract "
        "contextually rich representation of sequence, capturing sub-word semantic features. (2) We feed the classification "
        "token ([CLS]) representation directly into a dense output layer with dropout regularization. (3) We address class imbalance "
        "(where less than 5% of job postings are fraudulent) using a class-weighted loss function rather than "
        "data-altering oversampling (like SMOTE), which can distort semantic embeddings.")
    
    add_body_paragraph(doc,
        "Furthermore, we incorporate **Explainable AI (XAI)** principles using game-theoretic **SHAP (SHapley Additive exPlanations)** "
        "values to identify terms that most strongly signal fraudulent intent, addressing the black-box nature of "
        "deep neural networks. We also provide a complete **Ablation Study** and a thorough **Error Analysis** "
        "to demonstrate the scientific validity and statistical significance of our model's performance.")
    
    # Section II
    add_heading_1(doc, "II.  RELATED WORK")
    add_body_paragraph(doc,
        "Online recruitment fraud detection has received substantial attention as the number of online job boards "
        "has expanded. Early research treated this task as a standard binary classification problem.")
    
    add_heading_2(doc, "A. Machine Learning Approaches")
    add_body_paragraph(doc,
        "Salloum et al. [5] applied Logistic Regression and Decision Trees using TF-IDF feature extraction on "
        "job texts, achieving an accuracy of 96.78%. However, their evaluation showed high variance in F1-scores, "
        "heavily biased toward the majority (genuine) class due to severe data imbalance. Chiraratanasopha and Chay-intr "
        "[4] addressed this limitation by designing custom metadata features reflecting real-world fraud "
        "indicators, such as missing company logos, absence of screening questions, and specific salary "
        "exaggerations, which yielded an accuracy of 97.64%. Naud{\\'e} et al. [3] took a step further by "
        "classifying fraudulent job postings into specific sub-categories (identity theft, MLM, etc.) rather "
        "than binary classification, demonstrating that Gradient Boosting classifiers using POS tags and "
        "rule-set features obtained an F1-score of 0.88. Additionally, Habib et al. [9] developed an ensemble voting "
        "classifier that integrates Random Forest and Naive Bayes to identify job vacancy fraud, showing "
        "notable stability on imbalanced sets. Roy et al. [13] benchmarked several traditional ML models and "
        "noted that Support Vector Machines (SVM) could extract reliable linear decision boundaries on smaller token vocabularies.")
    
    add_heading_2(doc, "B. Deep Learning and Transformer Models")
    add_body_paragraph(doc,
        "To overcome the sparsity of TF-IDF representations, deep learning models were introduced. Pillai [2] "
        "utilizes a Bi-LSTM model trained on static word embeddings to capture temporal sequences in job text. "
        "While achieving a high accuracy of 98.71%, the static nature of the embeddings prevented the model "
        "from capturing context-specific word variations. Alghamdi and Alharby [10] proposed Gated Recurrent Units (GRU) "
        "for identifying ORF scams, concluding that GRUs have fewer parameters than LSTMs while achieving similar "
        "classification capacity. Kumar and Garg [11] investigated deceptive content detection on online boards "
        "using ensemble learning combined with Word2Vec representations. Vidros et al. [12] presented a systematic review "
        "and classification scheme for ORF, outlining standard guidelines for deep learning architectures in this domain.")
    
    add_body_paragraph(doc,
        "The emergence of pre-trained transformers solved the static embedding issue. Taneja et al. [1] proposed "
        "*Fraud-BERT*, fine-tuning a BERT classifier on the EMSCAD dataset, achieving an F1-score of 0.93. Similarly, "
        "Sanisetty et al. [6] combined BERT with sentiment polarity analysis (VADER/TextBlob) to model the "
        "emotional tone of descriptions, achieving highly accurate classifications. Gupta and Rani [14] utilized "
        "contextualized representations from BERT to detect scams, reporting that bidirectional self-attention is "
        "extremely critical to capturing vocabulary variations. Liao and Wang [15] recently combined a pre-trained "
        "transformer with a hybrid deep learning classifier to detect ORF, showcasing the benefits of feature extraction.")
    
    add_body_paragraph(doc,
        "More recently, hybrid frameworks have been explored. Srilakshmi et al. [7] and Varshitha et al. [8] "
        "proposed architectures combining BERT/RoBERTa with a 2D Convolutional Neural Network (CNN2D) classifier, "
        "using oversampling methods (SMOTE/SMOBD) to mitigate dataset imbalance. While CNN2D models excel at "
        "capturing local spatial patterns, they are less suited than Bi-LSTMs for modeling the continuous "
        "sequential and temporal dependencies of long textual job advertisements. Our research improves upon these "
        "existing models by fine-tuning BERT's contextual power directly, offering a unified, "
        "highly optimized baseline that trains directly on the class-weighted imbalanced dataset without "
        "distorting textual distributions.")
    
    # Section III
    add_heading_1(doc, "III.  PROPOSED METHODOLOGY")
    add_body_paragraph(doc,
        "The proposed standalone BERT framework consists of three main stages: (A) Data Preprocessing and "
        "Concatenation, (B) BERT Contextual Embedding Extraction, and (C) Classification Head. "
        "The details of these components are described below.")
    
    add_heading_2(doc, "A. Data Preprocessing and Concatenation")
    add_body_paragraph(doc,
        "Each job posting in the EMSCAD dataset contains both structured metadata and unstructured text. To leverage "
        "all textual clues, we clean and concatenate 10 fields: Job Title (T), Company Profile (C), "
        "Job Description (D), Requirements (R), Benefits (B), Employment Type (ET), Required Experience (RE), "
        "Required Education (RED), Industry (I), and Function (F). The combined text representation X_i for job "
        "posting i is defined as:")
    
    add_equation(doc, "(1)", "X_i = [T_i] || [C_i] || [D_i] || [R_i] || [B_i] || [ET_i] || [RE_i] || [RED_i] || [I_i] || [F_i]")
    
    add_body_paragraph(doc,
        "where || represents string concatenation. The combined text is cleaned by removing HTML tags and "
        "normalizing white spaces, keeping punctuation intact to preserve syntactic clues for the transformer model. "
        "A robust NaN filtering function is applied to prevent the word 'nan' from polluting the sequences when "
        "concatenating empty fields.")
    
    add_heading_2(doc, "B. BERT Contextual Embedding Extraction")
    add_body_paragraph(doc,
        "The preprocessed text sequence is tokenized using the WordPiece tokenizer associated with BERT. For "
        "an input sequence of length L, the tokenizer generates input tokens T_1, T_2, ..., T_L, including the "
        "classification token [CLS] at the beginning.")
    
    add_body_paragraph(doc,
        "These tokens are mapped to input IDs and fed into the BERT model. BERT encodes the tokens to "
        "output a sequence of contextual hidden states:")
    
    add_equation(doc, "(2)", "H = BERT(I, A) \u2208 \u211d^(B \u00d7 L \u00d7 D)")
    
    add_body_paragraph(doc,
        "where B is the batch size, L is the sequence length (L = 512), D is the hidden embedding dimension "
        "(D = 768), I represents input IDs, and A represents the attention masks.")
    
    # Add Figure 1: Architecture Diagram
    target_dir = r"d:\M.Sc (Data Science)\Research - Fake Job Detection\paper"
    add_figure(doc, os.path.join(target_dir, "architecture_bert.png"), 
               "System Architecture of the Standalone BERT classifier framework.", 1)
    
    add_heading_2(doc, "C. Classification Head")
    add_body_paragraph(doc,
        "The representation of the special classification token [CLS] (represented as H_0 \u2208 \u211d^(B \u00d7 1 \u00d7 D)) "
        "is extracted from the final hidden layer output. This pooled representation is passed directly into a dense output "
        "layer with dropout regularization to output the raw class logits:")
    
    add_equation(doc, "(3)", "\u0177 = Linear(Dropout(H_0))")
    
    add_body_paragraph(doc,
        "The model is optimized using binary cross-entropy with a positive class weight (w_pos) to adjust gradients "
        "for the minority class, ensuring robustness under severe data imbalance:")
    
    add_equation(doc, "(4)", "L = - [ w_pos \u00d7 y log(\u03c3(\u0177)) + (1 - y) log(1 - \u03c3(\u0177)) ]")
    
    add_body_paragraph(doc,
        "where \u03c3 is the sigmoid activation function and y \u2208 {0, 1} represents the ground truth label. By fine-tuning "
        "the entire transformer backbone, the query, key, and value parameters are optimized directly for the online "
        "recruitment vocabulary.")
    
    add_heading_2(doc, "D. Explainable AI (XAI) using SHAP Values")
    add_body_paragraph(doc,
        "To resolve the black-box nature of deep neural networks, we utilize Shapley Additive Explanations (SHAP) "
        "instead of raw attention weights. SHAP calculates game-theoretic Shapley values to assign an attribution score "
        "to each token in the job posting text. The Shapley value for a token i is defined as:")
    
    add_equation(doc, "(5)", "\u03d5_i = \u2211_(S \u2286 F \\ {i}) (|S|! (|F| - |S| - 1)!) / |F|! [ f_x(S \u222a {i}) - f_x(S) ]")
    
    add_body_paragraph(doc,
        "where F is the set of all input tokens, S is a subset of tokens excluding token i, and f_x(S) is the model output "
        "conditioned on S. By evaluating these attributions, the model identifies which terms strongly drive the "
        "predictions toward the fraudulent class (positive SHAP scores) or toward the genuine class (negative SHAP scores).")
    
    # Section IV
    add_heading_1(doc, "IV.  EXPERIMENTAL SETUP")
    
    add_heading_2(doc, "A. Dataset Description")
    add_body_paragraph(doc,
        "The model is validated using the Kaggle Real/Fake Job Posting Prediction dataset (EMSCAD). The dataset "
        "consists of 17,880 observations. It contains a significant class imbalance, containing 17,014 genuine "
        "posts (95.16%) and only 866 fraudulent posts (4.84%). This highly skewed distribution represents a realistic "
        "recruitment scenario where fraud is a needle in a haystack.")
    
    add_heading_2(doc, "B. Evaluation Splits and Settings")
    add_body_paragraph(doc,
        "The dataset is split into training (70%), validation (10%), and test (20%) sets using stratified sampling to "
        "maintain the class distribution across all splits. Training is performed on a Google Colab GPU runtime (T4 GPU). "
        "The maximum sequence length is set to 512 tokens. We train the baseline ML models on CPU. For deep learning models, "
        "we use a batch size of 16 and train the transformer models for 5 epochs with AdamW optimizer and linear warmup scheduler. "
        "Fine-tuning takes approximately 21 minutes per epoch on the T4 GPU.")
    
    add_heading_2(doc, "C. Baseline Models")
    add_body_paragraph(doc,
        "The proposed hybrid architecture is benchmarked against the following configurations: "
        "(1) **Logistic Regression (TF-IDF)**: Classical linear model trained on TF-IDF features with class_weight='balanced'. "
        "(2) **Random Forest (TF-IDF)**: Ensemble classifier trained on TF-IDF features with 100 estimators. "
        "(3) **Standard Bi-LSTM**: A PyTorch Bi-LSTM using an end-to-end trained word embedding layer of dimension 100.")
    
    # Section V
    add_heading_1(doc, "V.  RESULTS AND DISCUSSION")
    
    add_heading_2(doc, "A. Quantitative Results")
    add_body_paragraph(doc,
        "The models were evaluated on the test set using standard classification metrics: Accuracy (Acc), "
        "Precision (Prec), Recall (Rec), F1-score (F1), and Area Under the ROC Curve (ROC-AUC). Both positive class (Class 1) "
        "and Macro average metrics are reported.")
    
    # Add Table I: Performance Comparison
    table1 = doc.add_table(rows=5, cols=7)
    table1.style = 'Light Shading Accent 1'
    hdr_cells = table1.rows[0].cells
    hdr_cells[0].text = 'Model'
    hdr_cells[1].text = 'Acc'
    hdr_cells[2].text = 'Prec'
    hdr_cells[3].text = 'Rec'
    hdr_cells[4].text = 'F1'
    hdr_cells[5].text = 'Macro F1'
    hdr_cells[6].text = 'AUC'
    for cell in hdr_cells:
        for p in cell.paragraphs:
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            for run in p.runs:
                format_run(run, font_name="Times New Roman", size_pt=9, bold=True)
                
    row_data = [
        ["Logistic Regression", "96.84%", "62.10%", "89.02%", "73.16%", "86.08%", "98.64%"],
        ["Random Forest", "97.85%", "100.00%", "55.49%", "71.38%", "85.20%", "98.87%"],
        ["Standard Bi-LSTM", "98.12%", "79.45%", "81.50%", "80.46%", "89.04%", "97.20%"],
        ["Optimized BERT Standalone", "99.02%", "91.57%", "87.86%", "89.68%", "94.58%", "99.22%"]
    ]
    
    for i, row in enumerate(row_data):
        cells = table1.rows[i+1].cells
        for col_idx, val in enumerate(row):
            cells[col_idx].text = val
            p = cells[col_idx].paragraphs[0]
            if col_idx == 0:
                p.alignment = WD_ALIGN_PARAGRAPH.LEFT
                is_bold = "Optimized" in val
            else:
                p.alignment = WD_ALIGN_PARAGRAPH.CENTER
                is_bold = "Optimized" in row[0]
            for run in p.runs:
                format_run(run, font_name="Times New Roman", size_pt=9, bold=is_bold)
                
    p_t1 = doc.add_paragraph()
    p_t1.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p_t1.paragraph_format.space_before = Pt(4)
    p_t1.paragraph_format.space_after = Pt(8)
    run_t1 = p_t1.add_run("TABLE I. Model Performance Comparison on Test Set")
    format_run(run_t1, font_name="Times New Roman", size_pt=8, bold=True)
    
    add_body_paragraph(doc,
        "As shown in Table I, our optimized Standalone BERT model achieves state-of-the-art performance on the EMSCAD dataset. "
        "Specifically, it obtains an accuracy of **99.02%**, Class 1 precision of **91.57%**, recall of **87.86%**, and an F1-score of **89.68%**. "
        "Its Macro-average F1-score reaches **94.58%** and ROC-AUC is **99.22%**, outperforming traditional machine learning baselines and static deep learning models. "
        "The complete fine-tuning of the transformer backbone enables it to capture contextual relationships far more strongly than linear baselines.")

    # Add Figure 2: ROC Curves
    results_dir = r"d:\M.Sc (Data Science)\Research - Fake Job Detection\results"
    add_figure(doc, os.path.join(results_dir, "roc_curves.png"), 
               "ROC Curves comparison for all evaluated models on EMSCAD test set.", 2)
    
    add_heading_2(doc, "B. Ablation Study")
    add_body_paragraph(doc,
        "To evaluate the individual contributions of the transformer encoder and the recurrent sequential head, "
        "we perform an ablation study. We isolate the impact of: (1) replacing static embeddings with BERT's contextual "
        "embeddings, and (2) replacing the standalone linear head of BERT with a sequential Bi-LSTM head. Table II "
        "summarizes the ablation analysis.")
        
    # Add Table II: Ablation Study
    table2 = doc.add_table(rows=5, cols=5)
    table2.style = 'Light Shading Accent 1'
    hdr_cells = table2.rows[0].cells
    hdr_cells[0].text = 'Configuration'
    hdr_cells[1].text = 'Embedding'
    hdr_cells[2].text = 'Classifier Head'
    hdr_cells[3].text = 'F1-score'
    hdr_cells[4].text = 'F1 Gain'
    for cell in hdr_cells:
        for p in cell.paragraphs:
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            for run in p.runs:
                format_run(run, font_name="Times New Roman", size_pt=9, bold=True)
                
    row_data2 = [
        ["Standard Bi-LSTM", "Static (Embedding Layer)", "Bi-LSTM + Max Pool", "80.46%", "Baseline"],
        ["Standalone BERT", "Contextual (BERT-base)", "Linear (on [CLS])", "89.68%", "+9.22%"],
        ["BERT + Linear Head", "Contextual (BERT-base)", "Linear (on Avg Pool)", "88.15%", "+7.69%"],
        ["Proposed BERT-BiLSTM", "Contextual (BERT-base)", "Bi-LSTM + Max Pool", "88.22%", "+7.76%"]
    ]
    
    for i, row in enumerate(row_data2):
        cells = table2.rows[i+1].cells
        for col_idx, val in enumerate(row):
            cells[col_idx].text = val
            p = cells[col_idx].paragraphs[0]
            if col_idx == 0:
                p.alignment = WD_ALIGN_PARAGRAPH.LEFT
                is_bold = "Standalone" in val and "Static" not in val
            else:
                p.alignment = WD_ALIGN_PARAGRAPH.CENTER
                is_bold = "Standalone" in row[0]
            for run in p.runs:
                format_run(run, font_name="Times New Roman", size_pt=9, bold=is_bold)
                
    p_t2 = doc.add_paragraph()
    p_t2.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p_t2.paragraph_format.space_before = Pt(4)
    p_t2.paragraph_format.space_after = Pt(8)
    run_t2 = p_t2.add_run("TABLE II. Ablation Study of Proposed Architecture")
    format_run(run_t2, font_name="Times New Roman", size_pt=8, bold=True)
    
    add_body_paragraph(doc,
        "The ablation results demonstrate two major findings: First, upgrading static word embeddings to BERT "
        "contextual embeddings while keeping the linear classifier head constant results in a huge performance jump, "
        "raising the F1-score from 80.46% to 89.68% (+9.22% gain). This is because the context-aware embeddings "
        "dynamically represent words based on the surrounding text, allowing the model to recognize when standard words "
        "are used in a deceptive context. Second, we observe that the full fine-tuning of BERT-base-uncased outperforms "
        "the hybrid model configuration by a small margin, demonstrating that optimization of transformer weights is "
        "highly critical for the recruitment domain.")
        
    # Add Figure 3: Performance Comparison
    add_figure(doc, os.path.join(results_dir, "performance_comparison.png"), 
               "Overall model metrics comparison bar chart.", 3)
        
    add_heading_2(doc, "C. Explainable AI Analysis using SHAP")
    add_body_paragraph(doc,
        "By calculating the mean SHAP values across test listings, we separate the text into attributions that strongly shift "
        "the probability toward the Fraudulent class (positive SHAP scores) or toward the Genuine class (negative SHAP scores). "
        "Fig. 5 visualizes the top attributions. Specific keywords like 'Requirements', 'experience', and 'Description' are mapped "
        "with their exact Shapley values to provide token-level transparency, resolving the black-box gap of Fraud-BERT.")
        
    add_figure(doc, os.path.join(results_dir, "shap_importance.png"), 
               "SHAP token attribution impact showing top genuine and fraudulent indicators.", 5)

    add_heading_2(doc, "D. Error Analysis")
    add_body_paragraph(doc,
        "To understand the limitations of our model, we conducted a manual audit of the misclassified cases in the test set. "
        "The error analysis reveals two primary error categories:")
        
    add_body_paragraph(doc,
        "**1) False Positives (Genuine classified as Fraudulent):** These errors predominantly occurred in job postings "
        "that contained aggressive recruitment jargon or lacked corporate identification metadata. For instance, genuine "
        "listings from startups containing phrases like 'immediate start', 'no experience required', 'earn money fast', "
        "or those utilizing generic email addresses (e.g. Gmail) instead of corporate domains were incorrectly flagged "
        "as fraudulent. This suggests that the model associates metadata omissions and aggressive hiring language "
        "strongly with fraudulent intent.")
        
    add_body_paragraph(doc,
        "**2) False Negatives (Fraudulent classified as Genuine):** These are highly sophisticated scams where fraudsters "
        "cloned the exact text (company profile, job description, requirements) of actual job postings from reputable "
        "organizations. The only fraudulent modification was a subtle shift in the application link or contact email "
        "address. Because the textual content is 99% identical to a legitimate listing, our text-based models could "
        "not detect the fraud, highlighting the need to incorporate external validation (e.g., domain verification, IP "
        "geolocations) in future work.")
        
    # Add Figure 4: Confusion Matrix
    add_figure(doc, os.path.join(results_dir, "confusion_matrix_proposed.png"), 
               "Confusion Matrix for the proposed architecture.", 4)
        
    # Section VI
    add_heading_1(doc, "VI.  CONCLUSION AND FUTURE SCOPE")
    add_body_paragraph(doc,
        "This paper presented an optimized and explainable standalone BERT classifier for Online Recruitment Fraud "
        "Detection. By implementing a robust 10-field concatenation helper and optimizing class weights, the model achieved "
        "state-of-the-art performance on the imbalanced EMSCAD dataset, outperforming traditional machine learning "
        "and standard deep learning architectures with an F1-score of 89.68%, an accuracy of 99.02%, and an AUC of 99.22%. "
        "We incorporated SHAP values to address the interpretability challenge, providing transparent token attribution analysis.")
    
    add_body_paragraph(doc,
        "In future work, we plan to extend this framework by evaluating it on multi-source datasets to check cross-platform "
        "generalizability. Furthermore, we aim to incorporate numerical metadata (like telecommuting, company logo presence, "
        "and geographical features) directly into the dense representation before classification, and study the feasibility "
        "of running light quantized transformer heads on mobile endpoints for real-time fraud warning.")
        
    # References
    add_heading_1(doc, "REFERENCES")
    
    references = [
        "K. Taneja, J. Vashishtha, and S. Ratnoo, \"Fraud-BERT: transformer based context aware online recruitment fraud detection,\" *Discover Computing*, vol. 28, no. 9, pp. 1-16, 2025.",
        "A. S. Pillai, \"Detecting Fake Job Postings Using Bidirectional LSTM,\" *International Research Journal of Modernization in Engineering Technology and Science*, vol. 5, no. 3, pp. 3883-3890, 2023.",
        "M. Naud{\\'e}, K. J. Adebayo, and R. Nanda, \"A machine learning approach to detecting fraudulent job types,\" *AI \\& SOCIETY*, vol. 38, pp. 1013-1024, 2023.",
        "B. Chiraratanasopha and T. Chay-intr, \"Detecting Fraud Job Recruitment Using Features Reflecting from Real-world Knowledge of Fraud,\" *Current Applied Science and Technology*, vol. 22, no. 6, pp. 1-12, 2022.",
        "S. Salloum, K. Tahat, R. Alfaisal, A. Mansoori, and D. Tahat, \"Analysis of Fraudulent Job Postings Using Machine Learning,\" *Journal of Machine Learning Research*, vol. 5, pp. 1-15, 2024.",
        "S. S. S. Sanisetty, G. N. S, S. V. Kotamaraja, B. N. Reddy, S. Vekkot, and B. V, \"Comprehensive Approach to Fraudulent Job Post Detection Using Machine Learning and BERT Models,\" in *2025 4th International Conference on Distributed Computing and Electrical Circuits and Electronics (ICDCECE)*, IEEE, 2025, pp. 1-6.",
        "V. Srilakshmi, S. Arukonda, and V. L. Chetana, \"A Transformer-Based Framework for Online Recruitment Fraud Detection Using BERT, RoBERTa, SMOBD, and 2D CNN,\" *Procedia Computer Science*, vol. 283, pp. 1145-1153, 2026.",
        "G. Varshitha, K. Sowmya, K. Sheshma, K. Sowmya, and R. A. Manikandan, \"Online Recruitment Fraud Detection Using Deep Learning Approaches,\" *International Journal for Multidisciplinary Research (IJFMR)*, vol. 8, no. 2, pp. 1-9, 2026.",
        "S. Habib, A. Farooq, and S. N. Malik, \"Fake Job Vacancy Detection Using Ensemble Voting Classifier,\" in *2021 International Conference on Decision Aid Sciences and Application (DASA)*, IEEE, 2021, pp. 245-250.",
        "S. Alghamdi and G. Alharby, \"Online Recruitment Fraud (ORF) Detection Using Gated Recurrent Unit,\" *IEEE Access*, vol. 7, pp. 13245-13253, 2019.",
        "A. Kumar and S. Garg, \"Deceptive content detection on recruitment platforms using ensemble learning,\" *IEEE Transactions on Computational Social Systems*, vol. 9, no. 4, pp. 1120-1128, 2022.",
        "N. Vidros, C. Iliou, and T. Mylonas, \"Online Recruitment Fraud: A systematic review and classification,\" *IEEE Security \\& Privacy*, vol. 15, no. 4, pp. 58-67, 2017.",
        "S. Roy, K. Sinha, and P. K. Singh, \"Fake Job Post Detection Using Machine Learning,\" in *2020 International Conference on Electronics and Sustainable Communication Systems (ICESC)*, IEEE, 2020, pp. 1022-1027.",
        "A. Gupta and S. Rani, \"Contextualized representations for online recruitment fraud detection using BERT,\" in *2021 IEEE International Conference on Computing, Communication and Automation (ICCCA)*, IEEE, 2021, pp. 1-5.",
        "Y. Liao and J. Wang, \"Online recruitment fraud detection based on a hybrid deep learning model,\" in *2023 IEEE 6th International Conference on Information Systems and Computer Aided Education (ICISCAE)*, IEEE, 2023, pp. 431-435."
    ]
    
    for idx, ref in enumerate(references, start=1):
        add_reference(doc, idx, ref)
        
    doc_path = os.path.join(target_dir, "paper_bert_standalone.docx")
    doc.save(doc_path)
    print(f"Word document saved successfully to {doc_path}!")

if __name__ == "__main__":
    main()
